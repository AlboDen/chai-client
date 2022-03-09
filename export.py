from docx import Document
from docx import shared

def time_now():
    import datetime
    now = str(datetime.datetime.now())
    # 2021-08-06 22:45:02.522446
    day = now[8:10]
    month = now[5:7]
    year = now[0:4]
    return str(day +"."+ month +"."+ year)

class EXPORT():
    were_save = "test_1.doc"
    were_save_2 = "test_2.doc"
    were_save_3 = "test_3.doc"
    def Schyot_na_oplatu(nomer_schyota = "_____",
                         data_create = time_now(),
                         CompanyName = "___ ______________________",
                         INN = "___________________", KPP = "___________________",
                         ADDRES = "cтрана _________, город ____________, ул._________________ ___№_____, ",
                         Phone = "_  _ _ _  _ _ _ - _ _  _ _",
                         Oborudovanie = "__________________________________________________________________",
                         Ser_nomer = "___________________________",
                         No_intable = "_____",
                         Product_name = "___________________________",
                         skolko = "_____",
                         Edinitsy = "________",
                         Cena= "______ _",
                         Summa = "______ _",
                         Total = "_________________ __",
                         Total_sum = "_________________ __",
                         Total_sum_1 = "_________________ __"
                         ):
        document = Document("C:/Users/user/PycharmProjects/pythonProject3/DATABASE/shablons_of_document/Schyot na oplatu.docx")
        all_paras = document.paragraphs

        if "вставка в текст" == "вставка в текст":
            all_paras[1].text = ""
            print(nomer_schyota)
            print(data_create)
            sentence = all_paras[1].add_run(f"Счет на оплату №{nomer_schyota} от {data_create}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(18)
            sentence.font.bold = True


            all_paras[4].text = ""
            sentence = all_paras[4].add_run(f"Заказчик:            "
                                            f"{CompanyName},"
                                            f" ИНН {INN},"
                                            f" КПП {KPP},"
                                            f" {ADDRES},"
                                            f" тел.: {Phone}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[6].text = ""
            sentence = all_paras[6].add_run(f"Оборудование:  "
                                            f"{Oborudovanie},"
                                            f" серийный номер "
                                            f"{Ser_nomer}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[9].text = ""
            sentence = all_paras[9].add_run(f"Итого: {Total}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[11].text = ""
            sentence = all_paras[11].add_run(f"Итого к оплате: {Total_sum}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[12].text = ""
            sentence = all_paras[12].add_run(f"Всего к оплате: {Total_sum_1}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

        all_tables = document.tables
        if "вставка в таблицу" == "вставка в таблицу":
            all_tables[1].rows[1].cells[0].text = f"{No_intable}"
            paragraphs = all_tables[1].rows[1].cells[0].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[1].rows[1].cells[0].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[1].rows[1].cells[1].text = f"{Product_name}"
            paragraphs = all_tables[1].rows[1].cells[1].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[1].rows[1].cells[1].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[1].rows[1].cells[2].text = f"{skolko}"
            paragraphs = all_tables[1].rows[1].cells[2].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[1].rows[1].cells[2].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[1].rows[1].cells[3].text = f"{Edinitsy}"
            paragraphs = all_tables[1].rows[1].cells[3].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[1].rows[1].cells[3].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[1].rows[1].cells[4].text = f"{Cena}"
            paragraphs = all_tables[1].rows[1].cells[4].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[1].rows[1].cells[4].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[1].rows[1].cells[5].text = f"{Summa}"
            paragraphs = all_tables[1].rows[1].cells[5].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[1].rows[1].cells[5].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER


        document.save(EXPORT.were_save)

    def Act_ocenki_teh_sostoyaniya_oborudovaniya(imya_oborudovania = "_______________________________________",
                                                 Nomer_acta = "__________",
                                                 seriyniy_nomer = "____________________",
                                                 opisanie_neispravnosti = "______________________________________________",
                                                 srok_dostavki = "__.__.____",
                                                 kategoriya_remonta = "_______________________________",
                                                 No_intable = "____",
                                                 Product_name = "______________________________________",
                                                 skolko = "____",
                                                 Edinitsy = "________",
                                                 Cena = "______ __",
                                                 Summa = "______ __",
                                                 Stoyemost_bez_nds = "______ __",
                                                 new_oborudivanie = "______ __",
                                                 data_create = time_now()
                                                 ):
        document = Document("C:/Users/user/PycharmProjects/pythonProject3/DATABASE/shablons_of_document/Act_ocenki_teh_sostoyaniya_oborudovaniya.docx")

        all_paras = document.paragraphs
        if "вставка в текст" == "вставка в текст":
            all_paras[8].text = ""
            sentence = all_paras[8].add_run(f"Акт оценки технического состояния оборудования №{Nomer_acta}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(18)
            sentence.font.bold = True

            all_paras[9].text = ""
            sentence = all_paras[9].add_run(f"Настоящий акт составлен о том, что была проведена диагностика технического "
                                            f"состояния оборудования {imya_oborudovania} и выявлены следующие неисправности:")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[10].text = ""
            sentence = all_paras[10].add_run(f"Наименование оборудования: {imya_oborudovania}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[11].text = ""
            sentence = all_paras[11].add_run(f"Серийный номер: {seriyniy_nomer}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[12].text = ""
            sentence = all_paras[12].add_run(f"Описание неисправности: {opisanie_neispravnosti}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[13].text = ""
            sentence = all_paras[13].add_run(f"Срок поставки электронных компонентов: до {srok_dostavki}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[14].text = ""
            sentence = all_paras[14].add_run(f"Категория ремонта: {kategoriya_remonta}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[17].text = ""
            sentence = all_paras[17].add_run(f"Итого стоимость рабос без НДС: {Stoyemost_bez_nds}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = True

            all_paras[19].text = ""
            sentence = all_paras[19].add_run(f"*Стоимость нового оборудования {new_oborudivanie}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[20].text = ""
            sentence = all_paras[20].add_run(f"Сроки поставки электронных компонентов и их стоимость действительны в течение 5 дней.")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = True

            all_paras[26].text = ""
            sentence = all_paras[26].add_run(f"{data_create}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            all_paras[26].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        all_tables = document.tables
        if "вставка в таблицу" == "вставка в таблицу":
            all_tables[0].rows[1].cells[0].text = f"{No_intable}"
            paragraphs = all_tables[0].rows[1].cells[0].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[0].rows[1].cells[0].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[0].rows[1].cells[1].text = f"{Product_name}"
            paragraphs = all_tables[0].rows[1].cells[1].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[0].rows[1].cells[1].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[0].rows[1].cells[2].text = f"{skolko}"
            paragraphs = all_tables[0].rows[1].cells[2].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[0].rows[1].cells[2].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[0].rows[1].cells[3].text = f"{Edinitsy}"
            paragraphs = all_tables[0].rows[1].cells[3].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[0].rows[1].cells[3].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[0].rows[1].cells[4].text = f"{Cena}"
            paragraphs = all_tables[0].rows[1].cells[4].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[0].rows[1].cells[4].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[0].rows[1].cells[5].text = f"{Summa}"
            paragraphs = all_tables[0].rows[1].cells[5].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[0].rows[1].cells[5].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

        document.save(EXPORT.were_save_2)
    def Kvitantia_na_priyom(nomer_kvitancii = "___________",
                            data_zapolneniya = time_now(),
                            zacazchik = "_______________________",
                            oborudovaniye = "__________________________________",
                            ser_nom = "__________________",
                            kolichestvo = "_____",
                            opisaniye_neispravnosti = "_________________________________________",
                            komplectatia = "_____________________________________________________________________",
                            meh_povrejdeniya = "________________________________________________________",
                            ):
        document = Document("C:/Users/user/PycharmProjects/pythonProject3/DATABASE/shablons_of_document/Kvitantia_na_priyom_oborudovaniya.docx")

        all_paras = document.paragraphs
        if "вставка в текст" == "вставка в текст":
            all_paras[4].text = ""
            sentence = all_paras[4].add_run(f"Квитанция на приём оборудования в ремонт № {nomer_kvitancii}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(14)
            sentence.font.bold = True

            all_paras[6].text = ""
            sentence = all_paras[6].add_run(f"Заказчик: {zacazchik}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = False

            all_paras[9].text = ""
            sentence = all_paras[9].add_run(f"Описание неисправности (со слов заказчика): {opisaniye_neispravnosti}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = True

            all_paras[10].text = ""
            sentence = all_paras[10].add_run(f"Комплектация: {komplectatia}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = True

            all_paras[11].text = ""
            sentence = all_paras[11].add_run(f"Механические повреждения: {meh_povrejdeniya}")
            sentence.font.name = 'Arial'
            sentence.font.size = shared.Pt(11)
            sentence.font.bold = True

        all_tables = document.tables
        if "вставка в таблицу" == "вставка в таблицу":
            all_tables[0].rows[1].cells[0].text = f"{oborudovaniye}"
            paragraphs = all_tables[0].rows[1].cells[0].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[0].rows[1].cells[0].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[0].rows[1].cells[1].text = f"{ser_nom}"
            paragraphs = all_tables[0].rows[1].cells[1].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[0].rows[1].cells[1].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

            all_tables[0].rows[1].cells[2].text = f"{kolichestvo}"
            paragraphs = all_tables[0].rows[1].cells[2].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = 'Arial'
            font.size = shared.Pt(11)
            font.bold = False
            from docx.enum.text import WD_TAB_ALIGNMENT
            all_tables[0].rows[1].cells[2].paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER

        document.save(EXPORT.were_save_3)
if __name__ == "__main__":
    EXPORT.Kvitantia_na_priyom()
